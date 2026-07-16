# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
SCskill - 质量手册生成脚本（docx 处理库）
==========================================
本模块只负责 docx 模板处理（提取概览、应用修改），不直接调用 LLM。
LLM 调用由后端 routes.py 用项目已有的 langchain create_llm 完成，
这样自动跟随用户在前端选择的模型，无需 skill 自己处理 LLM 配置。

提供以下函数供后端调用：
  - find_template()              查找知识库手册分类下的模板
  - convert_doc_to_docx(path)    .doc 转 .docx
  - extract_template_overview(doc)  提取模板结构概览
  - format_overview_for_llm(overview)  概览格式化为文本
  - format_survey_for_llm(survey)  调研数据格式化为文本
  - build_llm_prompt(overview_text, survey_text, survey=None)  构造 LLM 提示词，survey 参数用于计算实施日期
  - parse_llm_modifications(llm_text)  解析 LLM 返回的修改方案 JSON
  - apply_modifications(doc, modifications)  应用修改方案到 docx
"""
import os
import sys
import json
import re
import subprocess
import tempfile
import shutil
from pathlib import Path
from datetime import datetime


def ensure_packages():
    """确保 python-docx 已安装"""
    import importlib
    try:
        importlib.import_module('docx')
    except ImportError:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install',
                               'python-docx', '--quiet'])


ensure_packages()

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
TEMPLATES_DIR = SKILL_ROOT / "templates"

# 备用内置模板（仅当知识库手册分类无文件时使用）
TEMPLATE_FILE = "IATF16949_quality_manual_template.docx"


# ===================================================================
# 1. 模板查找 & .doc → .docx 转换
# ===================================================================

def find_template(agent_id=None, documents_dir=None):
    """查找模板文件（三级查找）：
    1. 企业内部文件知识库（agent_{agent_id}/手册/）下的 .docx/.doc 文件（递归搜索子目录）
    2. 全质知识库（external_kb/体系文件/手册/全质手册模板/）下的 .docx/.doc 文件
    3. SCskill 内置模板

    Args:
        agent_id: 智能体ID（如 'dfmea-risk-agent'）
        documents_dir: 数据目录路径（如 Path('C:/beifen/QZAGENT/data/documents')）
                       如果不传，自动从 SKILL_ROOT 向上查找

    返回 (template_path, need_convert, template_source)
    template_source: 'internal' / 'external' / 'builtin'
    """
    # 确定 documents_dir 路径
    if documents_dir is None:
        # 从 SKILL_ROOT 向上查找 data/documents 目录
        # SKILL_ROOT = skills/SCskill/
        # project_root = skills/ 的上一级 = QZAGENT/
        candidate = SKILL_ROOT.parent.parent / "data" / "documents"
        if candidate.exists():
            documents_dir = candidate
        else:
            # 兜底：用 SKILL_ROOT.parent/data/documents（可能不存在）
            documents_dir = SKILL_ROOT.parent / "data" / "documents"
    else:
        documents_dir = Path(documents_dir)

    print(f"[INFO] find_template: documents_dir={documents_dir}, agent_id={agent_id}")

    # 1. 企业内部文件知识库（递归搜索 手册/ 下所有子目录）
    if agent_id:
        manual_dir = documents_dir / f"agent_{agent_id}" / "手册"
        print(f"[INFO] 查找企业内部文件: {manual_dir} (exists={manual_dir.exists()})")
        if manual_dir.exists():
            found_internal = False
            # 递归搜索所有 .docx 文件
            for root, dirs, files in os.walk(str(manual_dir)):
                for f in sorted(files):
                    if f.lower().endswith('.docx') and not f.startswith('~$'):
                        print(f"[INFO] 从企业内部文件知识库找到模板: {f} (路径: {root})")
                        return Path(root) / f, False, 'internal'
            # 再搜 .doc 文件
            for root, dirs, files in os.walk(str(manual_dir)):
                for f in sorted(files):
                    if f.lower().endswith('.doc') and not f.startswith('~$'):
                        print(f"[INFO] 从企业内部文件知识库找到 .doc 模板: {f} (路径: {root})")
                        return Path(root) / f, True, 'internal'
            print(f"[INFO] 企业内部文件知识库 手册/ 下未找到 .docx/.doc 文件（用户未上传模板）")

    # 2. 全质知识库（external_kb/体系文件/手册/全质手册模板/）
    ext_template_dir = documents_dir / "external_kb" / "体系文件" / "手册" / "全质手册模板"
    print(f"[INFO] 查找全质知识库: {ext_template_dir} (exists={ext_template_dir.exists()})")
    if ext_template_dir.exists():
        for f in sorted(os.listdir(str(ext_template_dir))):
            if f.lower().endswith('.docx') and not f.startswith('~$'):
                # [需求] 全质知识库模板只使用 AAA/aaa 命名的文件
                if 'AAA' not in f and 'aaa' not in f:
                    print(f"[INFO] 跳过非AAA模板: {f}")
                    continue
                print(f"[INFO] 从全质知识库找到模板: {f}")
                return ext_template_dir / f, False, 'external'
        for f in sorted(os.listdir(str(ext_template_dir))):
            if f.lower().endswith('.doc') and not f.startswith('~$'):
                # [需求] 全质知识库模板只使用 AAA/aaa 命名的文件
                if 'AAA' not in f and 'aaa' not in f:
                    print(f"[INFO] 跳过非AAA模板: {f}")
                    continue
                print(f"[INFO] 从全质知识库找到 .doc 模板: {f}")
                return ext_template_dir / f, True, 'external'

    # 3. SCskill 内置模板
    builtin = TEMPLATES_DIR / TEMPLATE_FILE
    if builtin.exists():
        print(f"[INFO] 使用 SCskill 内置模板: {TEMPLATE_FILE}")
        return builtin, False, 'builtin'

    return None, False, None


def find_all_templates(agent_id=None, documents_dir=None):
    """查找所有手册模板（支持多个文件，用户可能上传多个手册分册）

    查找逻辑：
    1. 企业内部文件 agent_{id}/手册/ → 递归搜索所有 .docx/.doc 文件
    2. 全质知识库 external_kb/体系文件/手册/全质手册模板/ → AAA/aaa 命名的文件
    3. 无内置模板（手册无通用内置模板）

    返回 list of dict:
    [
      {"path": Path, "need_convert": bool, "source": "internal"/"external",
       "filename": "文件名"},
      ...
    ]
    """
    if documents_dir is None:
        candidate = SKILL_ROOT.parent.parent / "data" / "documents"
        if candidate.exists():
            documents_dir = candidate
        else:
            documents_dir = SKILL_ROOT.parent / "data" / "documents"
    else:
        documents_dir = Path(documents_dir)

    print(f"[INFO] find_all_templates: documents_dir={documents_dir}, agent_id={agent_id}")

    # 1. 收集企业内部文件
    internal_templates = []
    if agent_id:
        manual_dir = documents_dir / f"agent_{agent_id}" / "手册"
        print(f"[INFO] 查找企业内部文件: {manual_dir} (exists={manual_dir.exists()})")
        if manual_dir.exists():
            for root, dirs, files in os.walk(str(manual_dir)):
                for f in sorted(files):
                    if f.startswith('~$'):
                        continue
                    ext = f.lower().rsplit('.', 1)[-1] if '.' in f else ''
                    if ext in ('docx', 'doc'):
                        internal_templates.append({
                            "path": Path(root) / f,
                            "need_convert": (ext == 'doc'),
                            "source": 'internal',
                            "filename": f
                        })
                        print(f"[INFO] 企业内部文件: {f} (source=internal)")

    # 2. 收集全质知识库
    ext_templates = []
    ext_template_dir = documents_dir / "external_kb" / "体系文件" / "手册" / "全质手册模板"
    print(f"[INFO] 查找全质知识库: {ext_template_dir} (exists={ext_template_dir.exists()})")
    if ext_template_dir.exists():
        for f in sorted(os.listdir(str(ext_template_dir))):
            if f.startswith('~$'):
                continue
            ext = f.lower().rsplit('.', 1)[-1] if '.' in f else ''
            if ext in ('docx', 'doc'):
                # [需求] 全质知识库模板只使用 AAA/aaa 命名的文件
                if 'AAA' not in f and 'aaa' not in f:
                    print(f"[INFO] 跳过非AAA模板: {f}")
                    continue
                ext_templates.append({
                    "path": ext_template_dir / f,
                    "need_convert": (ext == 'doc'),
                    "source": 'external',
                    "filename": f
                })
                print(f"[INFO] 全质知识库: {f} (source=external)")

    # [新需求] 总是返回 internal + external 全部模板
    # - internal：用户上传的手册，routes.py 会用 AI 修改后作为主文件返回下载
    # - external：全质知识库的手册，routes.py 会作为"参考文件"原样复制返回下载（不修改公司名）
    # 这样用户上传不全时，缺失的部分可以从全质知识库补全作为参考
    all_templates = internal_templates + ext_templates
    print(f"[INFO] 总计找到 {len(all_templates)} 个手册模板（企业内部 {len(internal_templates)} + 全质知识库 {len(ext_templates)}）")
    for t in all_templates:
        print(f"  - [{t['source']}] {t['filename']}")

    return all_templates


def _find_libreoffice_executables():
    """返回当前机器上真实存在的 LibreOffice 可执行文件路径。

    不直接尝试不存在的固定命令，也不依赖/终止其他请求正在使用的
    LibreOffice 进程。`shutil.which` 覆盖 PATH 安装，后面的候选路径覆盖
    Windows 和常见 Unix/macOS 安装位置。
    """
    candidates = []
    for command in ('soffice', 'libreoffice'):
        resolved = shutil.which(command)
        if resolved:
            candidates.append(resolved)

    if os.name == 'nt':
        program_roots = (
            os.environ.get('ProgramW6432'),
            os.environ.get('ProgramFiles'),
            os.environ.get('ProgramFiles(x86)'),
        )
        for root in program_roots:
            if root:
                candidates.append(str(Path(root) / 'LibreOffice' / 'program' / 'soffice.exe'))
    else:
        candidates.extend((
            '/usr/bin/soffice',
            '/usr/local/bin/soffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        ))

    executables = []
    seen = set()
    for candidate in candidates:
        try:
            path = Path(candidate)
            if not path.is_file():
                continue
            normalized = str(path.resolve())
        except (OSError, TypeError, ValueError):
            continue
        key = os.path.normcase(normalized)
        if key not in seen:
            seen.add(key)
            executables.append(normalized)
    return executables


def _convert_with_libreoffice(soffice_path, doc_path, tmp_dir, timeout=120):
    """使用一个独立 LibreOffice 配置目录进行一次转换。

    每个调用都有自己的 UserInstallation，故同时转换的请求不会争抢
    LibreOffice 的全局 profile/锁文件。超时仅结束本次 subprocess；绝不
    杀掉系统里其他用户或其他 worker 的 soffice 进程。
    """
    profile_dir = Path(tmp_dir) / 'libreoffice-profile'
    profile_dir.mkdir(parents=True, exist_ok=True)
    converted = Path(tmp_dir) / f'{Path(doc_path).stem}.docx'
    command = [
        str(soffice_path),
        f'-env:UserInstallation={profile_dir.resolve().as_uri()}',
        '--headless',
        '--nologo',
        '--nodefault',
        '--nofirststartwizard',
        '--norestore',
        '--convert-to', 'docx',
        '--outdir', str(tmp_dir),
        str(doc_path),
    ]
    try:
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        if result.returncode == 0 and converted.is_file() and converted.stat().st_size > 0:
            print('[INFO] LibreOffice 转换成功')
            return str(converted)
        detail = (result.stderr or result.stdout or '').strip().replace('\n', ' ')
        print(f'[WARN] LibreOffice ({soffice_path}) 转换失败，返回码 {result.returncode}: {detail[:200]}')
    except subprocess.TimeoutExpired:
        print(f'[WARN] LibreOffice ({soffice_path}) 转换超时（{timeout}s），将尝试 Word COM（如可用）')
    except Exception as e:
        print(f'[WARN] LibreOffice ({soffice_path}) 转换失败: {e}')
    finally:
        # 输出文件在 tmp_dir 根目录；独立 profile 可以立即清理，避免锁/缓存积累。
        shutil.rmtree(profile_dir, ignore_errors=True)
    return None


def _convert_with_word_com(doc_path, tmp_dir):
    """Windows Word COM 兜底转换；无论成功失败都清理 COM 对象。"""
    pythoncom = None
    word = None
    source_doc = None
    initialized = False
    try:
        import pythoncom as _pythoncom
        import win32com.client

        pythoncom = _pythoncom
        pythoncom.CoInitialize()
        initialized = True
        # DispatchEx 创建独立实例，避免关闭/影响用户正在使用的 Word 窗口。
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        source_doc = word.Documents.Open(str(doc_path), ReadOnly=True, AddToRecentFiles=False)
        out = Path(tmp_dir) / 'converted.docx'
        source_doc.SaveAs2(str(out), FileFormat=16)
        if out.is_file() and out.stat().st_size > 0:
            print('[INFO] Word COM 转换成功')
            return str(out)
        print('[WARN] Word COM 转换未生成有效 .docx 文件')
    except Exception as e:
        print(f'[WARN] Word COM 转换失败: {e}')
    finally:
        if source_doc is not None:
            try:
                source_doc.Close(SaveChanges=0)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if initialized and pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
    return None


def convert_doc_to_docx(doc_path):
    """把 .doc 转为临时 .docx，成功返回路径、失败返回 ``None``。

    转换使用每次请求独立的 LibreOffice 用户配置，避免并发请求互相抢锁。
    失败的临时目录会立即清理；成功时保留目录，使调用方可以继续读取返回文件。
    """
    source = Path(doc_path)
    tmp_dir = Path(tempfile.mkdtemp(prefix='doc2docx_'))
    succeeded = False
    try:
        if not source.is_file():
            print(f'[WARN] .doc 文件不存在，无法转换: {source}')
            return None

        soffice_paths = _find_libreoffice_executables()
        for soffice_path in soffice_paths:
            converted = _convert_with_libreoffice(soffice_path, source, tmp_dir)
            if converted:
                succeeded = True
                return converted

        if not soffice_paths:
            print('[WARN] 未找到可用的 LibreOffice 可执行文件，将尝试 Word COM')
        converted = _convert_with_word_com(source, tmp_dir)
        if converted:
            succeeded = True
            return converted
        return None
    finally:
        if not succeeded:
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ===================================================================
# 2. 提取模板结构概览（带索引的纯文本）
# ===================================================================

def extract_template_overview(doc, max_paras=None):
    """提取模板的结构化概览，用于喂给 LLM。
    返回 dict：
      {
        "paragraphs": [{"index": N, "style": "...", "text": "..."}, ...],
        "tables":     [{"index": T, "rows": [[cell_text,...],...]}, ...],
        "headers":    [{"section": S, "text": "..."}, ...],
        "footers":    [{"section": S, "text": "..."}, ...]
      }
    """
    overview = {"paragraphs": [], "tables": [], "headers": [], "footers": []}

    # 段落（只保留非空的）
    for i, p in enumerate(doc.paragraphs):
        if max_paras and i >= max_paras:
            break
        text = p.text.strip()
        if not text:
            continue
        try:
            style = p.style.name if p.style else ''
        except Exception:
            style = ''
        # 跳过目录条目（避免污染 LLM 视野）
        if 'toc' in (style or '').lower():
            continue
        overview["paragraphs"].append({
            "index": i,
            "style": style,
            "text": text[:500]  # 截断超长段落
        })

    # 表格
    for ti, t in enumerate(doc.tables):
        rows = []
        for row in t.rows:
            cells = []
            for c in row.cells:
                cell_text = c.text.strip().replace('\n', ' | ')
                cells.append(cell_text[:200])
            rows.append(cells)
        overview["tables"].append({"index": ti, "rows": rows})

    # 页眉/页脚
    for si, sec in enumerate(doc.sections):
        for hf_name in ['header', 'first_page_header', 'even_page_header']:
            hf = getattr(sec, hf_name, None)
            if not hf:
                continue
            texts = []
            for p in hf.paragraphs:
                if p.text.strip():
                    texts.append(p.text.strip())
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if c.text.strip():
                            texts.append(c.text.strip().replace('\n', ' | '))
            if texts:
                overview["headers"].append({
                    "section": si,
                    "type": hf_name,
                    "text": " | ".join(texts)[:500]
                })
        for hf_name in ['footer', 'first_page_footer', 'even_page_footer']:
            hf = getattr(sec, hf_name, None)
            if not hf:
                continue
            texts = []
            for p in hf.paragraphs:
                if p.text.strip():
                    texts.append(p.text.strip())
            # 页脚中的编号、公司名、版次常放在表格单元格内。若只提取段落，
            # LLM 看不到它们，便无法提出相应 table_cell/global_replace 修改。
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if c.text.strip():
                            texts.append(c.text.strip().replace('\n', ' | '))
            if texts:
                overview["footers"].append({
                    "section": si,
                    "type": hf_name,
                    "text": " | ".join(texts)[:500]
                })

    return overview


def format_overview_for_llm(overview):
    """把概览格式化成 LLM 易读的文本"""
    lines = []
    lines.append("=== 模板段落（非空，已跳过目录） ===")
    for p in overview["paragraphs"]:
        lines.append(f"[P{p['index']}] ({p['style']}) {p['text']}")
    lines.append("")
    lines.append("=== 模板表格 ===")
    for t in overview["tables"]:
        lines.append(f"--- Table {t['index']} ({len(t['rows'])} rows) ---")
        for ri, row in enumerate(t['rows']):
            lines.append(f"  T{t['index']}.R{ri}: {row}")
    lines.append("")
    lines.append("=== 页眉 ===")
    for h in overview["headers"]:
        lines.append(f"[H{h['section']}.{h['type']}] {h['text']}")
    lines.append("")
    lines.append("=== 页脚 ===")
    for f in overview["footers"]:
        lines.append(f"[F{f['section']}.{f['type']}] {f['text']}")
    return "\n".join(lines)


def format_survey_for_llm(survey):
    """把调研数据格式化成 LLM 易读的文本（含字段说明）"""
    field_labels = {
        'sv_company_name': '公司名称',
        'sv_cert_other': '其他证书',
        'sv_chairman': '董事长',
        'sv_legal_rep': '法人代表',
        'sv_gm': '总经理',
        'sv_deputy_gm': '副总经理',
        'sv_mgmt_rep': '管理者代表',
        'sv_leader_group_leader': '贯标组长',
        'sv_leader_group_members': '贯标组员',
        'sv_iso_office_head': '贯标办主任',
        'sv_iso_office_members': '贯标办成员',
        'sv_auditors': '内审员',
        'sv_products': '体系覆盖产品',
        'sv_process_flow': '生产流程',
        'sv_location': '地理位置',
        'sv_area': '占地面积',
        'sv_building_area': '建筑面积',
        'sv_staff_total': '正式员工人数',
        'sv_staff_mgmt': '管理技术人员数',
        'sv_staff_edu': '学历分布',
        'sv_equipment': '设备情况',
        'sv_customers': '主要客户',
        'sv_address': '公司地址',
        'sv_contact': '联系人',
        'sv_phone': '电话',
        'sv_fax': '传真',
        'sv_mobile': '手机',
        'sv_purpose': '公司宗旨/经营理念',
        'sv_quality_policy': '质量方针',
        'sv_quality_goal': '质量目标',
        'sv_cert_date': '认证日期',
        'sv_audit_date': '审核日期',
        'sv_rest_day': '休息日',
        'sv_design_dev': '有无设计开发',
        'sv_filler_name': '填写人',
        'sv_filler_phone': '填写人手机',
        'sv_certs': '已有证书',
        'sv_org': '机构设置',
        'sv_org_custom_rows': '自定义部门',
    }
    lines = ["=== 用户体系调研数据 ==="]
    for key, label in field_labels.items():
        val = survey.get(key, '')
        if isinstance(val, (list, dict)):
            val = json.dumps(val, ensure_ascii=False)
        if val:
            lines.append(f"{label}（{key}）：{val}")
        else:
            lines.append(f"{label}（{key}）：[未填写]")
    return "\n".join(lines)


# ===================================================================
# 3. 构造 LLM 提示词
# ===================================================================

def _parse_date(date_str):
    """将各种格式的日期字符串解析为 datetime 对象，失败返回 None。
    支持: 2024年3月15日, 2024-03-15, 2024/03/15, 2024.03.15, ISO格式"""
    if not date_str or not isinstance(date_str, str):
        return None
    date_str = date_str.strip()
    if not date_str:
        return None
    # 尝试各种格式
    formats = [
        '%Y年%m月%d日', '%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d',
        '%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M:%S.%f',
        '%Y-%m-%d %H:%M:%S',
        '%Y年%m月%d日 %H:%M:%S',
    ]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    # 尝试用正则提取数字
    m = re.search(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', date_str)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except ValueError:
            return None
    return None


def _contains_16949(certs):
    """判断证书列表中是否包含 IATF 16949"""
    if not certs:
        return False
    if isinstance(certs, str):
        return '16949' in certs
    if isinstance(certs, (list, tuple)):
        for c in certs:
            if isinstance(c, str) and '16949' in c:
                return True
    return False


def _subtract_months(dt, months):
    """从日期减去指定月数，返回新的 datetime。自动处理月末溢出（如 3月31日 - 1月 = 2月28/29日）。"""
    import calendar
    total_months = dt.year * 12 + dt.month - 1 - months
    new_year = total_months // 12
    new_month = total_months % 12 + 1
    max_day = calendar.monthrange(new_year, new_month)[1]
    new_day = min(dt.day, max_day)
    return dt.replace(year=new_year, month=new_month, day=new_day)


def calculate_implementation_date(survey):
    """根据证书类型和第一阶段开始时间，计算手册实施日期（发布/生效日期）。

    逻辑：
    - 证书含 IATF 16949:
      - 未填写第一阶段开始时间 → 计划取得认证证书日期 - 18个月
      - 已填写第一阶段开始时间 → 第一阶段开始时间 - 15个月
    - 非 16949（如 ISO 9001）:
      - 未填写第一阶段开始时间 → 计划取得认证证书日期 - 5个月
      - 已填写第一阶段开始时间 → 第一阶段开始时间 - 4个月

    返回 (implementation_date_str, description_str)
    implementation_date_str 格式: "2024年3月15日"
    """
    certs = survey.get('sv_certs', [])
    is_16949 = _contains_16949(certs)

    audit_date_str = (survey.get('sv_audit_date') or '').strip()
    cert_date_str = (survey.get('sv_cert_date') or '').strip()

    audit_date = _parse_date(audit_date_str)
    cert_date = _parse_date(cert_date_str)

    if is_16949:
        if audit_date:
            impl_date = _subtract_months(audit_date, 15)
            desc = f"IATF 16949认证，第一阶段开始时间={audit_date_str}，向前推15个月作为手册实施日期"
        elif cert_date:
            impl_date = _subtract_months(cert_date, 18)
            desc = f"IATF 16949认证，计划取得认证证书日期={cert_date_str}，向前推18个月作为手册实施日期"
        else:
            today = datetime.now()
            impl_date = today
            desc = "IATF 16949认证，但认证日期和第一阶段开始时间均未填写，回退使用当前日期"
    else:
        if audit_date:
            impl_date = _subtract_months(audit_date, 4)
            desc = f"非16949认证，第一阶段开始时间={audit_date_str}，向前推4个月作为手册实施日期"
        elif cert_date:
            impl_date = _subtract_months(cert_date, 5)
            desc = f"非16949认证，计划取得认证证书日期={cert_date_str}，向前推5个月作为手册实施日期"
        else:
            today = datetime.now()
            impl_date = today
            desc = "非16949认证，但认证日期和第一阶段开始时间均未填写，回退使用当前日期"

    impl_date_str = f"{impl_date.year}年{impl_date.month}月{impl_date.day}日"
    return impl_date_str, desc


def build_llm_prompt(overview_text, survey_text, survey=None):
    """构造 LLM 提示词，返回 (system_prompt, user_prompt)
    使用 NDJSON（每行一个 JSON 对象）格式，方便流式增量解析。

    Args:
        overview_text: 模板结构概览文本
        survey_text: 调研数据格式化文本
        survey: 原始调研数据 dict（用于计算实施日期，可选）
    """
    today = datetime.now()
    today_str = f"{today.year}年{today.month}月{today.day}日"
    year_str = str(today.year)

    # 计算实施日期：根据证书类型和第一阶段开始时间倒推
    if survey:
        impl_date_str, impl_date_desc = calculate_implementation_date(survey)
    else:
        impl_date_str = today_str
        impl_date_desc = "未提供调研原始数据，使用当前日期"

    system = (
        "你是质量手册智能生成助手。你会收到一份质量手册模板的结构概览（带段落索引P#、表格T#.R#、页眉H#、页脚F#）"
        "和用户填写的体系调研数据。你的任务是根据调研数据，决定模板中哪些位置需要修改，逐条输出修改方案。\n\n"
        "【输出格式 - 极其重要】\n"
        "每条修改方案单独输出为一行 JSON 对象（NDJSON 格式），不要包裹在数组里，不要输出任何其他文字。"
        "每生成一条就立即输出一行，不要等所有方案想完再一起输出。输出完所有方案后，最后一行写：===END===\n\n"
        "示例输出（每行一个 JSON，逐行输出）：\n"
        '{"type":"global_replace","old":"山东AAA机械制造有限公司","new":"诸暨正和金属有限公司","reason":"整体替换公司名"}\n'
        '{"type":"header_replace","old":"AAA-QM-2021","new":"正和-QM-' + year_str + '","reason":"替换文件编号"}\n'
        '{"type":"paragraph","index":81,"new_text":" 总经理：王大明    （签名）","reason":"颁布令填入总经理姓名"}\n'
        '{"type":"paragraph","index":92,"new_text":"    诸暨正和金属有限公司位于...","reason":"重写公司简介"}\n'
        '{"type":"table_cell","table":0,"row":2,"col":3,"new_text":"王大明","reason":"批准人=总经理"}\n'
        '===END===\n\n'
        "修改类型说明：\n"
        "- paragraph: 把段落 P#index 整段替换为 new_text（保留段落格式）\n"
        "- table_cell: 把表格 T#table 第 row 行第 col 列单元格替换为 new_text\n"
        "- global_replace: 在全文（段落+表格+页眉页脚）中把 old 替换为 new（用于公司名整体替换）\n"
        "- header_replace: 仅在页眉页脚中把 old 替换为 new\n\n"
        "关键规则：\n"
        "1. 公司名必须整体替换：模板里出现的所有原公司名变体（如\"山东AAA机械制造有限公司\"、\"AAA\"、"
        "\"纳赫\"等）都要整体替换为用户填写的公司名。绝不能产生\"诸暨正和金属有限公司机械制造有限公司\"这种错误拼接。\n"
        "2. 人名填入对应签字位置：总经理姓名填入颁布令\"总经理：\"行和文件控制表\"批准\"栏；"
        "管理者代表姓名填入文件控制表\"审核\"栏；贯标办主任填入\"制订\"栏；董事长姓名如有签字位也填入。\n"
        "3. 质量方针：找到模板的\"质量方针\"章节（通常是 2.1），把原方针正文段落替换为用户填写的质量方针。\n"
        "4. 质量目标：找到模板的\"质量目标\"章节（通常是 2.3），把原目标正文段落替换为用户填写的质量目标。\n"
        "5. 公司简介：找到模板的\"公司简介\"章节（通常是 1.1），根据调研数据（公司名+地点+地址+产品+规模+设备+客户+宗旨）"
        "重写 1-3 段简介正文。\n"
        "6. 公司联络：地址、电话、传真按字段替换到\"公司联络\"章节对应行。\n"
        "7. 公司宗旨/经营理念：替换模板里\"公司经营理念：...\"段。\n"
        "8. 文件编号：模板里的文件编号（如 AAA-QM-2021）替换为公司简称+当前年份（简称取公司名前4字）。\n"
        "9. 实施日期（重要——已按认证准备周期倒推计算）：手册应先在体系运行前发布实施，"
        "根据证书类型和第一阶段开始时间倒推得出实施日期为 " + impl_date_str + "。\n"
        "   【计算依据】" + impl_date_desc + "。\n"
        "   模板里所有\"X年X月X日起实施\"、\"发布日期\"、\"生效日期\"、\"实施日期\"等表达手册生效时间的日期，"
        "都要替换为 " + impl_date_str + "。\n"
        "   注意：\"计划取得认证证书日期\"不要改，那是认证目标日期，不是手册的日期。\n"
        "   手册的日期是体系开始运行的日期，必须早于认证日期。\n"
        "10. 不要修改 IATF16949/ISO9001 标准条款内容、不要修改程序文件引用名（如《文件管理程序》）。\n"
        "11. 如果调研数据某字段为空，跳过对应修改。\n"
        "12. 修改方案要全面，覆盖所有该改的位置（公司名通常在封面、页眉、承诺书、简介等处多次出现）。\n"
        "13. new_text 中的换行用 \\n 转义，不要在 JSON 字符串中放真实换行符。\n"
        "14. 想到一条就立即输出一条，不要等想完所有再输出。\n"
        "15. 【重要】所有页面都要检查，不能因为内容少就跳过：\n"
        "    - 修改记录页：文件编号要替换（如 AAA-QM-2021 → 公司简称-QM-年份）\n"
        "    - 职能分配表：公司名、部门名称要替换为用户调研数据中的部门\n"
        "    - 封面页、附件页、任命书等：公司名、人名、日期、编号都要检查\n"
        "    - 空白表格：表头里的公司名、文件编号也要替换\n"
        "16. 【重要】表格内的内容必须检查：\n"
        "    - 表格单元格里的公司名、文件编号、人名、日期都要用 table_cell 替换\n"
        "    - 职能分配表的部门名称要与用户调研数据中的组织架构一致\n"
        "    - 修改记录表的文件编号要替换\n"
        "17. 【重要】即使页面内容很少（只有表头或几个字），也要检查是否有公司名/编号/人名/日期需要替换。\n"
    )

    user = (
        f"当前日期：{today_str}\n\n"
        f"{survey_text}\n\n"
        f"{overview_text}\n\n"
        "请分析以上模板和调研数据，逐条输出 NDJSON 格式的修改方案，每条一行，最后输出 ===END===。"
    )

    return system, user


def build_diff_prompt(ext_overview_text, user_manual_text, ext_filename):
    """构造「差异提取」LLM 提示词。

    让 AI 对比全质模板概览 vs 用户上传手册，输出「全质模板有但用户手册未覆盖」的章节。

    输出格式（每行一条，以「差异章节:」开头）：
        差异章节: 5.3 内部审核 - 全质模板要求每月一次内部审核，用户手册未提及审核频次...
        差异章节: 7.2 培训 - 全质模板包含培训记录表要求，用户手册未包含...
        ===END===

    如果用户手册已完全涵盖全质模板内容，只输出 ===END=== 不输出差异行。
    """
    system = (
        "你是质量手册内容对比专家。你会收到两份文档：\n"
        "1. 全质知识库的质量手册模板（结构概览，含段落和表格）\n"
        "2. 用户上传的一个或多个手册（结构概览）\n\n"
        "你的任务是：找出「全质模板中包含、但用户手册未覆盖」的章节或条款。\n\n"
        "【判定原则】\n"
        "- 如果全质模板有某个章节/条款，用户手册中有对应内容（即使表述不同），算「已覆盖」\n"
        "- 如果全质模板有某个章节/条款，用户手册中根本没有相关内容，算「未覆盖」（差异）\n"
        "- 如果用户手册只是表述更简略但主题一致，算「已覆盖」\n"
        "- 如果用户手册的内容更丰富（额外内容），不算差异（差异只关心全质模板有、用户没有的部分）\n\n"
        "【输出格式 - 极其重要】\n"
        "每条差异单独输出一行，以「差异章节:」开头，格式：\n"
        "差异章节: <章节标题或编号> - <全质模板中的内容摘要，包含关键要点>\n\n"
        "示例：\n"
        "差异章节: 5.3 内部审核 - 全质模板要求制定年度审核计划、每月一次内审、审核报告管理评审输入，用户手册未提及审核频次和报告管理\n"
        "差异章节: 7.2 培训 - 全质模板包含培训需求识别、培训计划、培训记录表、培训效果评估，用户手册未包含培训记录表要求\n"
        "===END===\n\n"
        "重要：\n"
        "1. 只输出差异，不要输出已覆盖的章节\n"
        "2. 如果没有差异（用户手册已完全涵盖），直接输出 ===END===\n"
        "3. 每条差异要包含「全质模板中的具体内容要点」，让用户能从这一条记录中看懂缺失了什么\n"
        "4. 不要输出其他任何说明文字\n"
    )

    user = (
        f"全质知识库模板：{ext_filename}\n\n"
        f"=== 全质模板结构概览 ===\n{ext_overview_text}\n\n"
        f"=== 用户上传手册结构概览 ===\n{user_manual_text}\n\n"
        "请对比以上两份文档，逐行输出差异章节，最后输出 ===END===。"
    )

    return system, user




def build_covered_prompt(ext_overview_text, user_manual_text, ext_filename):
    """构造「已覆盖章节识别」LLM 提示词。

    让 AI 对比全质模板概览 vs 用户上传手册，输出全质模板中已被用户手册
    覆盖的段落和表格索引。routes.py 据此从参考文件中删除已覆盖部分，保留
    真正缺失的内容。后端兼容旧版 ``indices`` 字段，但新输出使用更明确的
    ``paragraph_indices`` / ``table_indices``。
    """
    system = (
        "你是质量手册内容对比专家。你会收到两份文档：\n"
        "1. 全质知识库的质量手册模板（结构概览，每段带 P# 索引）\n"
        "2. 用户上传的一个或多个手册（结构概览）\n\n"
        "你的任务是：找出「全质模板中已被用户手册覆盖」的段落和表格索引，输出给后端用于删除。\n\n"
        "【判定原则】\n"
        "- 如果全质模板某段落的内容在用户手册中有对应主题（即使表述不同），算「已覆盖」\n"
        "- 如果用户手册只是更简略但主题一致，算「已覆盖」\n"
        "- 如果用户手册完全没相关内容，算「未覆盖」（保留为差异）\n"
        "- 段落索引必须严格基于全质模板概览中的 P# 编号\n"
        "- 表格索引必须严格基于全质模板概览中的 Table # 编号；表格主题、职责矩阵、编号等已被用户文件覆盖时，也要标记该表格\n\n"
        "【输出格式 - 极其重要】\n"
        "每条输出一行 JSON 对象（NDJSON），格式：\n"
        '{"covered": true, "paragraph_indices": [P索引1, ...], "table_indices": [表格索引1, ...], "reason": "简要说明"}\n\n'
        "示例：\n"
        '{"covered": true, "paragraph_indices": [3, 4, 5], "table_indices": [], "reason": "5.3 内部审核 - 用户手册已覆盖"}\n'
        '{"covered": true, "paragraph_indices": [], "table_indices": [1], "reason": "职责分配表已由用户手册覆盖"}\n'
        "===END===\n\n"
        "重要：\n"
        "1. paragraph_indices 和 table_indices 必须是全质模板概览中真实存在的索引；没有时填 []\n"
        "2. 一条 JSON 可以包含多个连续段落或多个表格索引（同一章节）\n"
        "3. 不要把整个文档都标记为已覆盖，除非用户手册真的涵盖了所有内容\n"
        "4. 标题段落（如 '5.3 内部审核' 单独成段）也要包含在 paragraph_indices 中\n"
        "5. 不要输出其他任何说明文字，只输出 JSON 行 + ===END===\n"
        "6. 如果用户手册完全没覆盖任何内容，直接输出 ===END===（不输出 JSON 行）\n"
        "7. 旧字段 indices 仅供后端兼容历史响应；本次新响应不要使用它\n"
    )

    user = (
        f"全质知识库模板：{ext_filename}\n\n"
        f"=== 全质模板结构概览 ===\n{ext_overview_text}\n\n"
        f"=== 用户上传手册结构概览 ===\n{user_manual_text}\n\n"
        "请对比以上两份文档，逐行输出已被用户手册覆盖的段落索引和表格索引（NDJSON 格式），最后输出 ===END===。"
    )

    return system, user

def parse_ndjson_line(line):
    """解析一行 NDJSON 为 dict，失败返回 None。
    支持去掉 markdown 代码块标记。"""
    line = line.strip()
    if not line or line == '===END===':
        return None
    # 去掉 markdown 代码块标记
    if line.startswith('```'):
        line = line.lstrip('`')
        # 可能是 ```json 或 ```
        line = line.replace('json', '', 1).strip()
    if line.endswith('```'):
        line = line[:-3].strip()
    if not line.startswith('{') or not line.endswith('}'):
        return None
    try:
        return json.loads(line)
    except json.JSONDecodeError:
        # 尝试修复尾随逗号
        fixed = re.sub(r',\s*([}\]])', r'\1', line)
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            return None


# ===================================================================
# 4. 解析 LLM 修改方案
# ===================================================================

def parse_llm_modifications(llm_text):
    """从 LLM 响应中解析出 modifications 列表"""
    text = llm_text.strip()
    text = re.sub(r'^```(?:json)?\s*', '', text)
    text = re.sub(r'\s*```\s*$', '', text)

    m = re.search(r'\{[\s\S]*\}', text)
    if not m:
        print(f"[WARN] LLM 响应中找不到 JSON")
        return []

    try:
        data = json.loads(m.group(0))
    except json.JSONDecodeError:
        # 尝试修复常见 JSON 错误（尾随逗号）
        fixed = re.sub(r',\s*([}\]])', r'\1', m.group(0))
        try:
            data = json.loads(fixed)
        except json.JSONDecodeError as e2:
            print(f"[WARN] JSON 解析失败: {e2}")
            return []

    mods = data.get('modifications', [])
    if not isinstance(mods, list):
        return []
    return mods


# ===================================================================
# 5. 应用修改方案到 docx
# ===================================================================

def set_paragraph_text(p, new_text):
    """保留段落第一个 run 的格式，把整段文本设为 new_text，其余 run 清空"""
    if new_text is None:
        new_text = ''
    if not p.runs:
        p.add_run(str(new_text))
        return
    p.runs[0].text = str(new_text)
    for r in p.runs[1:]:
        r.text = ''


def _paragraph_runs_in_document_order(paragraph):
    """取得段落中所有文本 run（包括超链接内的 run）。"""
    try:
        # Paragraph.runs 不包含 hyperlink 内的 run。直接从 XML 取得它们，
        # 以免一次全局替换把超链接或相邻文字合并到第一个 run。
        from docx.text.run import Run
        return [Run(element, paragraph) for element in paragraph._p.xpath('.//w:r')]
    except Exception:
        # 极少数非标准文档可退回 python-docx 的公开 API；仍然不会执行
        # 旧实现那种“把整段塞入第一个 run”的破坏性替换。
        return list(paragraph.runs)


def replace_text_in_paragraph(p, old, new):
    """跨 run 替换文本，同时保留未命中部分原有的 run 格式。

    新文本继承命中字符串首字符所在 run 的格式；其余字符仍留在原 run。
    因此仅替换一个公司名不会抹掉同段的粗体、斜体、超链接等格式。
    """
    if old is None:
        return False
    old = str(old)
    new = '' if new is None else str(new)
    if not old:
        return False

    runs = _paragraph_runs_in_document_order(p)
    if not runs:
        return False
    run_texts = [run.text or '' for run in runs]
    full_text = ''.join(run_texts)
    if old not in full_text:
        return False

    # 为原始字符串的每个字符记录其所属 run。替换后，未命中的字符仍回到
    # 原 run；替换文本放入命中起点的 run，继承最接近的局部格式。
    tokens = []
    for run_index, text in enumerate(run_texts):
        tokens.extend((char, run_index) for char in text)

    result_tokens = []
    cursor = 0
    while True:
        match_start = full_text.find(old, cursor)
        if match_start < 0:
            result_tokens.extend(tokens[cursor:])
            break
        result_tokens.extend(tokens[cursor:match_start])
        replacement_run_index = tokens[match_start][1]
        result_tokens.extend((char, replacement_run_index) for char in new)
        cursor = match_start + len(old)

    replacement_texts = ['' for _ in runs]
    for char, run_index in result_tokens:
        replacement_texts[run_index] += char

    for run, old_text, replacement_text in zip(runs, run_texts, replacement_texts):
        if old_text != replacement_text:
            run.text = replacement_text
    return True


def replace_text_in_cell(cell, old, new):
    """在表格单元格内做整体替换"""
    changed = False
    for p in cell.paragraphs:
        if replace_text_in_paragraph(p, old, new):
            changed = True
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                if replace_text_in_cell(c, old, new):
                    changed = True
    return changed


def apply_global_replace(doc, old, new):
    """在全文（段落+表格+页眉页脚）做替换"""
    if old is None:
        return 0
    old = str(old)
    new = '' if new is None else str(new)
    if not old or old == new:
        return 0
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            if replace_text_in_paragraph(p, old, new):
                count += 1
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if old in c.text:
                    if replace_text_in_cell(c, old, new):
                        count += 1
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer]:
            if not hf:
                continue
            for p in hf.paragraphs:
                if old in p.text:
                    if replace_text_in_paragraph(p, old, new):
                        count += 1
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if old in c.text:
                            if replace_text_in_cell(c, old, new):
                                count += 1
    return count


def apply_header_replace(doc, old, new):
    """仅在页眉页脚做替换"""
    if old is None:
        return 0
    old = str(old)
    new = '' if new is None else str(new)
    if not old or old == new:
        return 0
    count = 0
    for sec in doc.sections:
        for hf in [sec.header, sec.first_page_header, sec.even_page_header,
                   sec.footer, sec.first_page_footer, sec.even_page_footer]:
            if not hf:
                continue
            for p in hf.paragraphs:
                if old in p.text:
                    if replace_text_in_paragraph(p, old, new):
                        count += 1
            for t in hf.tables:
                for row in t.rows:
                    for c in row.cells:
                        if old in c.text:
                            if replace_text_in_cell(c, old, new):
                                count += 1
    return count


def apply_paragraph_replace(doc, index, new_text):
    """把指定索引的段落整段替换"""
    if index < 0 or index >= len(doc.paragraphs):
        print(f"[WARN] 段落索引 {index} 越界（总段落数 {len(doc.paragraphs)}）")
        return False
    p = doc.paragraphs[index]
    # 【节分界段落保护】如果段落含 sectPr（节分界符），禁止修改内容
    # 节分界段落必须保持空白，填入内容会导致节属性错乱、页眉引用丢失、空白页
    from docx.oxml.ns import qn
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
        print(f"[INFO] 跳过节分界段落 P{index}（含 sectPr，禁止修改）")
        return False
    set_paragraph_text(p, new_text)
    return True


def delete_paragraphs_by_indices(doc, indices):
    """从 docx 中删除指定索引的段落（用于生成补缺参考文件时删除用户已覆盖的章节）。

    注意：
    - 从大到小排序后逐个删除，避免索引移位
    - 跳过节分界段落（含 sectPr），避免破坏文档结构
    - 段落被删除后，其所在表格不会被删除（表格是独立元素）

    Args:
        doc: python-docx Document 对象
        indices: 要删除的段落索引列表

    Returns:
        int: 实际删除的段落数
    """
    from docx.oxml.ns import qn
    if not indices:
        return 0
    # [Bug 修复] 先转 int 再过滤负数，避免字符串索引崩溃
    _parsed = []
    for i in indices:
        try:
            _parsed.append(int(i))
        except (ValueError, TypeError):
            continue
    unique_indices = sorted(set(i for i in _parsed if i >= 0), reverse=True)
    deleted = 0
    total = len(doc.paragraphs)
    for idx in unique_indices:
        if idx >= total:
            continue
        p = doc.paragraphs[idx]
        # 跳过节分界段落
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            print(f"[INFO] 删除时跳过节分界段落 P{idx}")
            continue
        # 从父元素中移除该段落
        p._element.getparent().remove(p._element)
        deleted += 1
    print(f"[INFO] 删除了 {deleted} 个段落（请求 {len(unique_indices)} 个，含越界/节分界跳过）")
    return deleted


def delete_tables_by_indices(doc, indices):
    """安全删除文档正文中指定索引的顶级表格。

    索引与 :func:`extract_template_overview` 的 ``Table #`` 一致。按降序删除
    可避免后续索引移动；非法、重复和越界索引会被忽略。页眉/页脚表格不属于
    ``doc.tables``，不会被这个面向参考文件正文的函数误删。
    """
    if not indices:
        return 0

    parsed_indices = []
    for index in indices:
        try:
            parsed_indices.append(int(index))
        except (TypeError, ValueError):
            continue
    unique_indices = sorted(set(index for index in parsed_indices if index >= 0), reverse=True)
    if not unique_indices:
        return 0

    deleted = 0
    total = len(doc.tables)
    for index in unique_indices:
        if index >= total:
            print(f"[WARN] 删除表格时索引 {index} 越界（总表格数 {total}）")
            continue
        table = doc.tables[index]
        parent = table._element.getparent()
        if parent is None:
            print(f"[WARN] 删除表格时找不到父节点 T{index}，已跳过")
            continue
        try:
            parent.remove(table._element)
            deleted += 1
        except (AttributeError, ValueError) as exc:
            print(f"[WARN] 删除表格 T{index} 失败: {exc}")
    print(f"[INFO] 删除了 {deleted} 个表格（请求 {len(unique_indices)} 个，含越界/无效跳过）")
    return deleted


def apply_table_cell_replace(doc, table_idx, row_idx, col_idx, new_text):
    """把指定表格的指定单元格替换"""
    if table_idx < 0 or table_idx >= len(doc.tables):
        print(f"[WARN] 表格索引 {table_idx} 越界（总表格数 {len(doc.tables)}）")
        return False
    t = doc.tables[table_idx]
    if row_idx < 0 or row_idx >= len(t.rows):
        print(f"[WARN] 行索引 {row_idx} 越界（表格 {table_idx} 共 {len(t.rows)} 行）")
        return False
    row = t.rows[row_idx]
    if col_idx < 0 or col_idx >= len(row.cells):
        print(f"[WARN] 列索引 {col_idx} 越界（表格 {table_idx} 行 {row_idx} 共 {len(row.cells)} 列）")
        return False
    cell = row.cells[col_idx]
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], new_text)
        for p in cell.paragraphs[1:]:
            for r in list(p.runs):
                r.text = ''
    else:
        cell.text = str(new_text)
    return True


def _order_overlapping_replacements(modifications):
    """确保互相包含的全局/页眉替换从最长旧值开始执行。

    LLM 有时会同时给出 ``AAA -> 新公司`` 和
    ``AAA企业 -> 新公司有限公司``。如果短值先执行，后者会再也匹配不到，
    留下“新公司企业”这种残缺文本。仅重排旧值彼此包含的替换组，其他修改
    保持原顺序，以兼顾正确性与原有行为。
    """
    planned = list(modifications)
    replacement_types = {'global_replace', 'header_replace'}
    positions = [
        index for index, mod in enumerate(planned)
        if isinstance(mod, dict) and mod.get('type') in replacement_types
        and mod.get('old') not in (None, '')
    ]
    if len(positions) < 2:
        return planned

    old_values = {
        index: str(planned[index].get('old'))
        for index in positions
    }
    # 找到“旧值相互包含”的连通组；只有这些组才需要按长度重排。
    neighbours = {index: set() for index in positions}
    for offset, left in enumerate(positions):
        for right in positions[offset + 1:]:
            left_old = old_values[left]
            right_old = old_values[right]
            if left_old in right_old or right_old in left_old:
                neighbours[left].add(right)
                neighbours[right].add(left)

    visited = set()
    reordered = False
    for start in positions:
        if start in visited:
            continue
        component = []
        pending = [start]
        visited.add(start)
        while pending:
            current = pending.pop()
            component.append(current)
            for neighbour in neighbours[current]:
                if neighbour not in visited:
                    visited.add(neighbour)
                    pending.append(neighbour)
        if len(component) < 2:
            continue

        target_positions = sorted(component)
        source_positions = sorted(
            component,
            key=lambda index: (-len(old_values[index]), index),
        )
        source_modifications = [planned[index] for index in source_positions]
        for target, source_modification in zip(target_positions, source_modifications):
            if planned[target] is not source_modification:
                reordered = True
            planned[target] = source_modification

    if reordered:
        print('[INFO] 已按旧值长度重排重叠的全局/页眉替换，避免短名称破坏长名称')
    return planned


def apply_modifications(doc, modifications):
    """应用所有修改方案"""
    stats = {
        'paragraph': 0,
        'table_cell': 0,
        'global_replace': 0,
        'header_replace': 0,
        'unknown': 0,
        'failed': 0,
    }
    for i, mod in enumerate(_order_overlapping_replacements(modifications)):
        try:
            mod_type = mod.get('type', '')
            reason = (mod.get('reason', '') or '')[:50]
            if mod_type == 'paragraph':
                idx = int(mod.get('index', -1))
                new_text = mod.get('new_text', '')
                if apply_paragraph_replace(doc, idx, new_text):
                    stats['paragraph'] += 1
                    print(f"  [P{idx}] OK {reason}")
                else:
                    stats['failed'] += 1
            elif mod_type == 'table_cell':
                ti = int(mod.get('table', -1))
                ri = int(mod.get('row', -1))
                ci = int(mod.get('col', -1))
                new_text = mod.get('new_text', '')
                if apply_table_cell_replace(doc, ti, ri, ci, new_text):
                    stats['table_cell'] += 1
                    print(f"  [T{ti}.R{ri}.C{ci}] OK {reason}")
                else:
                    stats['failed'] += 1
            elif mod_type == 'global_replace':
                old = mod.get('old', '')
                new = mod.get('new', '')
                n = apply_global_replace(doc, old, new)
                stats['global_replace'] += n
                print(f"  [G] '{old}' -> '{new}' ({n} 处) OK {reason}")
            elif mod_type == 'header_replace':
                old = mod.get('old', '')
                new = mod.get('new', '')
                n = apply_header_replace(doc, old, new)
                stats['header_replace'] += n
                print(f"  [H] '{old}' -> '{new}' ({n} 处) OK {reason}")
            else:
                stats['unknown'] += 1
                print(f"  [?] 未知类型: {mod_type}")
        except Exception as e:
            stats['failed'] += 1
            print(f"  [ERROR] 修改 #{i} 失败: {e}")
    return stats


def remove_even_page_headers_footers(doc):
    """兼容旧调用的安全 no-op，保留奇偶页页眉/页脚和分节结构。

    奇偶页页眉/页脚是 Word 的合法布局，不能把它当成 LibreOffice 转换
    产生的异常而删除；移动疑似“空白页”的分节符同样会破坏真实模板。
    保留此函数和 ``int`` 返回值是为了兼容历史调用方。
    """
    print('[INFO] 保留偶数页页眉页脚和分节结构（不再执行破坏性清理）')
    return 0

